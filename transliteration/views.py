import os
import io
import csv
import json
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from django.shortcuts import render
from .forms import UploadFileForm
from .utils import cyrillic_to_latin


def upload_file_view(request):
    form = UploadFileForm()

    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            f = request.FILES['file']
            filename = f.name
            file_type = os.path.splitext(filename)[1].lower()

            buffer = io.BytesIO()
            response = HttpResponse(content_type="application/octet-stream")
            response["Content-Disposition"] = f"attachment; filename=File{file_type}"

            if file_type == ".txt":
                raw_text = f.read().decode('utf-8')
                converted = cyrillic_to_latin(raw_text)
                response.write(converted)

            elif file_type == ".docx":
                old_doc = Document(f)
                new_doc = Document()

                for para in old_doc.paragraphs:
                    new_para = new_doc.add_paragraph()
                    new_para.alignment = para.alignment

                    for run in para.runs:
                        new_run = new_para.add_run(cyrillic_to_latin(run.text))
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name

                new_doc.save(buffer)
                buffer.seek(0)
                response = HttpResponse(
                    buffer.read(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                response["Content-Disposition"] = "attachment; filename=File.docx"

            elif file_type == ".csv":
                df = pd.read_csv(f)
                raw_text = df.to_csv(index=False)
                converted = cyrillic_to_latin(raw_text)
                output = io.StringIO()
                pd.read_csv(io.StringIO(converted)).to_csv(output, index=False)
                response.write(output.getvalue())

            elif file_type == ".html":
                soup = BeautifulSoup(f.read(), 'html.parser')
                for tag in soup.find_all(text=True):
                    tag.replace_with(cyrillic_to_latin(tag))
                response.write(str(soup))

            elif file_type == ".json":
                data = json.load(f)
                raw_text = json.dumps(data, ensure_ascii=False, indent=2)
                converted = cyrillic_to_latin(raw_text)
                response.write(converted)

            else:
                return HttpResponse("Unsupported file type!", status=400)

            return response

    return render(request, "transliteration/file_upload.html", {"form": form})